import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _add_bottom_border(paragraph, color="CCCCCC"):
    """Adds a thin bottom border to a paragraph — used under section headings."""
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _looks_like_heading(line: str) -> bool:
    return line.startswith("##") or (line.isupper() and len(line) < 40) or (len(line) < 40 and line.endswith(":"))


def generate_resume_docx(resume_text: str) -> bytes:
    doc = Document()

    # Narrow margins for resume
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    base_style = doc.styles["Normal"]
    base_style.font.name = "Calibri"
    base_style.font.size = Pt(10.5)

    lines = [line.rstrip() for line in (resume_text or "").strip().split("\n")]

    name_written = False
    contact_written = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()
            continue

        # First non-empty line — candidate name, large and centered
        if not name_written:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = True
            run.font.size = Pt(18)
            name_written = True
            continue

        # Next line with contact separators — email / phone / links, centered and muted
        if not contact_written and any(ch in stripped for ch in ("@", "|", "·")):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            contact_written = True
            continue

        # Section headings — bold, uppercase, underlined with a thin rule
        if _looks_like_heading(stripped):
            heading_text = stripped.replace("##", "").replace(":", "").strip().upper()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            _add_bottom_border(p)
            continue

        # Bullet points
        if stripped.startswith(("•", "-", "*")):
            bullet_text = stripped.lstrip("•-* ").strip()
            p = doc.add_paragraph(bullet_text, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(10.5)
            continue

        # Everything else — regular body text (job titles, dates, plain lines)
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.font.size = Pt(10.5)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()